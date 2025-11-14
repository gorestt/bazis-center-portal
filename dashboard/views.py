
import os
from django.shortcuts import render, get_object_or_404, redirect
from django.contrib.auth.decorators import login_required
from django.core.paginator import Paginator
from django.http import JsonResponse, FileResponse, HttpResponseForbidden
from django.utils import timezone
from django.conf import settings
from django.contrib.auth.models import User

from .models import OrderQueue, KPIRecord, Incident, Shift, Document, Report
from .forms import OrderForm, DocumentForm, ReportForm

def get_role(user):
    profile = getattr(user, 'profile', None)
    if profile:
        return profile.role
    if user.is_superuser:
        return 'admin'
    return 'client'

def ensure_sample_data():
    # пользователи для связей
    admin = User.objects.filter(username='admin').first()
    manager = User.objects.filter(username='manager').first()
    client = User.objects.filter(username='client').first()

    # Очередь заявок
    if OrderQueue.objects.count() == 0 and client:
        OrderQueue.objects.create(
            title='Недоступен портал клиентов',
            description='Пользователи сообщают о недоступности портала авторизации.',
            initiator=client,
            executor=manager or admin,
            status='in_progress',
            priority='high',
            sla_deadline=timezone.now() + timezone.timedelta(hours=4),
        )
        OrderQueue.objects.create(
            title='Ошибка при формировании отчёта',
            description='При генерации отчёта за месяц появляется сообщение об ошибке.',
            initiator=client,
            executor=manager or admin,
            status='new',
            priority='medium',
            sla_deadline=timezone.now() + timezone.timedelta(hours=8),
        )
        OrderQueue.objects.create(
            title='Уточнение прав доступа',
            description='Необходимо выдать права на просмотр отчётов для нового сотрудника.',
            initiator=client,
            executor=manager or admin,
            status='done',
            priority='low',
            sla_deadline=timezone.now() - timezone.timedelta(hours=1),
        )

    # KPI
    if KPIRecord.objects.count() == 0:
        base = timezone.now()
        for i in range(7):
            day = base - timezone.timedelta(days=i)
            KPIRecord.objects.create(
                metric='Среднее время решения',
                value=4.0 - i * 0.2,
                timestamp=day,
                service_name='Портал клиентов',
            )
            KPIRecord.objects.create(
                metric='Доступность сервиса',
                value=99.0 + i * 0.1,
                timestamp=day,
                service_name='Система биллинга',
            )

    # Инциденты
    if Incident.objects.count() == 0:
        Incident.objects.create(
            title='Снижение скорости обработки запросов',
            description='Зафиксировано увеличение времени отклика портала клиентов.',
            status='В работе',
            criticality='medium',
            detected_at=timezone.now() - timezone.timedelta(hours=3),
            related_order=OrderQueue.objects.first(),
        )
        Incident.objects.create(
            title='Кратковременная недоступность биллинга',
            description='Пользователи не могли формировать счета в течение 10 минут.',
            status='Закрыт',
            criticality='high',
            detected_at=timezone.now() - timezone.timedelta(days=1, hours=2),
            closed_at=timezone.now() - timezone.timedelta(days=1),
        )

    # Смены
    if Shift.objects.count() == 0 and (admin or manager):
        base_date = timezone.now().date()
        for i in range(5):
            Shift.objects.create(
                employee=manager or admin,
                date=base_date + timezone.timedelta(days=i),
                shift='day',
                comment='Плановая дневная смена',
                phone='+7 (900) 000-00-01',
            )
        Shift.objects.create(
            employee=admin or manager,
            date=base_date,
            shift='night',
            comment='Ночная смена дежурного инженера',
            phone='+7 (900) 000-00-02',
        )

    # Документы
    if Document.objects.count() == 0:
        docs_dir = settings.MEDIA_ROOT / 'docs'
        os.makedirs(docs_dir, exist_ok=True)
        try:
            from docx import Document as DocxDocument
            def make_doc(filename, title_text, body_text):
                doc = DocxDocument()
                doc.add_heading(title_text, level=1)
                doc.add_paragraph(body_text)
                full_path = docs_dir / filename
                doc.save(full_path)
            make_doc('reglament_incidents.docx', 'Регламент обработки инцидентов',
                     'Документ описывает порядок регистрации, классификации и эскалации инцидентов.')
            make_doc('reglament_shifts.docx', 'Регламент организации смен',
                     'Документ фиксирует правила формирования графика смен и обязанности дежурного персонала.')
            make_doc('instruction_portal.docx', 'Инструкция пользователя портала',
                     'Инструкция по работе с порталом и отслеживанию статусов обращений.')
        except Exception:
            # резервный вариант: текстовые файлы
            def make_txt(filename, body_text):
                full_path = docs_dir / filename
                with open(full_path, 'w', encoding='utf8') as f:
                    f.write(body_text)
            make_txt('reglament_incidents.txt', 'Регламент обработки инцидентов.')
            make_txt('reglament_shifts.txt', 'Регламент организации смен.')
            make_txt('instruction_portal.txt', 'Инструкция пользователя портала.')

        # создать записи в БД
        for fname, title, desc in [
            ('reglament_incidents.docx', 'Регламент обработки инцидентов', 'Порядок регистрации и сопровождения инцидентов.'),
            ('reglament_shifts.docx', 'Регламент организации смен', 'Описание процедуры планирования смен и ответственности.'),
            ('instruction_portal.docx', 'Инструкция пользователя портала', 'Руководство по работе с порталом для клиентов.'),
        ]:
            # если docx не создан, пробуем txt
            path_docx = docs_dir / fname
            path_txt = docs_dir / fname.replace('.docx', '.txt')
            if path_docx.exists():
                rel = f'docs/{fname}'
            elif path_txt.exists():
                rel = f'docs/{path_txt.name}'
            else:
                continue
            Document.objects.get_or_create(
                slug=fname.split('.')[0],
                defaults={
                    'title': title,
                    'description': desc,
                    'file': rel,
                    'access': 'public',
                }
            )

@login_required
def home(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role == 'client':
        return redirect('dashboard:client_home')
    open_orders = OrderQueue.objects.filter(status__in=['new','in_progress']).count()
    incidents_open = Incident.objects.exclude(status='Закрыт').count()
    kpi_count = KPIRecord.objects.count()
    return render(request, 'dashboard/home.html', {
        'open_orders': open_orders,
        'incidents_open': incidents_open,
        'kpi_count': kpi_count,
    })

@login_required
def queue_list(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')
    qs = OrderQueue.objects.select_related('initiator','executor').order_by('-created_at')
    status = request.GET.get('status')
    priority = request.GET.get('priority')
    if status:
        qs = qs.filter(status=status)
    if priority:
        qs = qs.filter(priority=priority)
    paginator = Paginator(qs, 20)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    return render(request, 'dashboard/queue_list.html', {
        'page_obj': page_obj,
        'status': status,
        'priority': priority,
    })

@login_required
def queue_create(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager', 'client']:
        return HttpResponseForbidden('Доступ запрещён.')
    if request.method == 'POST':
        form = OrderForm(request.POST)
        if form.is_valid():
            order = form.save(commit=False)
            order.initiator = request.user
            order.save()
            return redirect('dashboard:queue_list')
    else:
        form = OrderForm()
    return render(request, 'dashboard/queue_form.html', {'form': form, 'mode': 'create'})

@login_required
def queue_edit(request, pk):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')
    order = get_object_or_404(OrderQueue, pk=pk)
    if request.method == 'POST':
        form = OrderForm(request.POST, instance=order)
        if form.is_valid():
            form.save()
            return redirect('dashboard:queue_list')
    else:
        form = OrderForm(instance=order)
    return render(request, 'dashboard/queue_form.html', {'form': form, 'mode': 'edit', 'order': order})

@login_required
def queue_detail(request, pk):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager', 'client']:
        return HttpResponseForbidden('Доступ запрещён.')
    order = get_object_or_404(OrderQueue, pk=pk)
    # клиент видит только свои заявки
    if role == 'client' and order.initiator != request.user:
        return HttpResponseForbidden('Доступ запрещён.')
    return render(request, 'dashboard/queue_detail.html', {'order': order})

@login_required
def kpi_dashboard(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')
    last_30 = timezone.now() - timezone.timedelta(days=30)
    kpi = KPIRecord.objects.filter(timestamp__gte=last_30)
    series = {}
    for rec in kpi:
        series.setdefault(rec.metric, []).append({
            'timestamp': rec.timestamp.isoformat(),
            'value': rec.value,
        })
    return render(request, 'dashboard/kpi_dashboard.html', {
        'kpi': kpi,
        'series_json': series,
    })

@login_required
def incidents_list(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')
    items = Incident.objects.order_by('-detected_at')
    paginator = Paginator(items, 20)
    page_obj = paginator.get_page(request.GET.get('page'))
    return render(request, 'dashboard/incidents_list.html', {'page_obj': page_obj})

@login_required
def shifts_list(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')
    items = Shift.objects.select_related('employee').order_by('date')
    return render(request, 'dashboard/shifts_list.html', {'items': items})

@login_required
def reports_panel(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')
    if request.method == 'POST':
        form = ReportForm(request.POST)
        if form.is_valid():
            report = form.save(commit=False)
            report.author = request.user
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument()
                doc.add_heading('Отчёт по сервисам', level=1)
                doc.add_paragraph(f"Тип: {report.get_report_type_display()}")
                doc.add_paragraph(f"Период: {report.period_from} – {report.period_to}")
                path_dir = settings.MEDIA_ROOT / 'reports'
                os.makedirs(path_dir, exist_ok=True)
                filename = f"report_{report.report_type}_{report.period_from}_{report.period_to}.docx"
                full_path = path_dir / filename
                doc.save(full_path)
                report.file.name = f"reports/{filename}"
            except Exception:
                path_dir = settings.MEDIA_ROOT / 'reports'
                os.makedirs(path_dir, exist_ok=True)
                filename = f"report_{report.report_type}_{report.period_from}_{report.period_to}.txt"
                full_path = path_dir / filename
                with open(full_path, 'w', encoding='utf8') as f:
                    f.write('Отчёт по сервисам\n')
                    f.write(f"Тип: {report.get_report_type_display()}\n")
                    f.write(f"Период: {report.period_from} – {report.period_to}\n")
                report.file.name = f"reports/{filename}"
            report.save()
            return redirect('dashboard:reports_panel')
    else:
        form = ReportForm()
    reports = Report.objects.order_by('-created_at')
    return render(request, 'dashboard/reports_panel.html', {'reports': reports, 'form': form})

@login_required
def report_download(request, pk):
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')
    report = get_object_or_404(Report, pk=pk)
    if not report.file:
        raise HttpResponseForbidden('Файл отчёта не найден.')
    full_path = report.file.path
    return FileResponse(open(full_path, 'rb'), as_attachment=True, filename=os.path.basename(full_path))

@login_required
def docs_manage(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')
    if request.method == 'POST':
        form = DocumentForm(request.POST, request.FILES)
        if form.is_valid():
            form.save()
            return redirect('dashboard:docs_manage')
    else:
        form = DocumentForm()
    docs = Document.objects.all()
    return render(request, 'dashboard/docs_manage.html', {'docs': docs, 'form': form})

@login_required
def client_home(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role != 'client':
        return redirect('dashboard:home')
    my_orders = OrderQueue.objects.filter(initiator=request.user).order_by('-created_at')
    return render(request, 'dashboard/client_home.html', {'orders': my_orders})

# API views
@login_required
def queue_api(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')
    qs = OrderQueue.objects.order_by('-created_at')
    status = request.GET.get('status')
    priority = request.GET.get('priority')
    if status:
        qs = qs.filter(status=status)
    if priority:
        qs = qs.filter(priority=priority)
    data = [{
        'id': o.id,
        'title': o.title,
        'status': o.status,
        'priority': o.priority,
        'created_at': o.created_at.isoformat(),
    } for o in qs[:200]]
    return JsonResponse({'results': data})

@login_required
def kpi_api(request):
    ensure_sample_data()
    role = get_role(request.user)
    if role not in ['admin', 'manager']:
        return HttpResponseForbidden('Доступ запрещён.')
    metric = request.GET.get('metric')
    qs = KPIRecord.objects.all()
    if metric:
        qs = qs.filter(metric=metric)
    data = [{
        'metric': r.metric,
        'value': r.value,
        'timestamp': r.timestamp.isoformat(),
        'service_name': r.service_name,
    } for r in qs[:500]]
    return JsonResponse({'results': data})
